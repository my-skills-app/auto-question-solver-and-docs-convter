"""Export solved paper to JSON + Word (.docx) using CSV template fields."""

from __future__ import annotations

import csv
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .csv_template import CSV_HEADERS, questions_to_rows
from .json_schema import processing_report


def _strip_html(text: str) -> str:
    text = text or ""
    text = re.sub(r"(?is)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</p\s*>", "\n", text)
    text = re.sub(r"(?is)<[^>]+>", "", text)
    return (
        text.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&nbsp;", " ")
        .strip()
    )


def save_json(doc: dict[str, Any], path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(doc, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    return path


def save_report(doc: dict[str, Any], path: Path) -> Path:
    report = processing_report(doc)
    report["generated_at"] = datetime.now().isoformat(timespec="seconds")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    return path


def _rows_from_doc(
    doc: dict[str, Any], *, set_name: str = "Paper Name"
) -> list[dict[str, str]]:
    """Prefer bilingual raw questions; fall back to normalized schema."""
    questions = doc.get("questions_raw") or doc.get("questions") or []
    if questions and (
        questions[0].get("question_en")
        or questions[0].get("question_hi")
        or questions[0].get("option1_en")
    ):
        return questions_to_rows(questions, set_name=set_name)

    # Rebuild from normalized schema
    rebuilt = []
    for q in questions:
        opts = q.get("options") or {}
        letters = ["A", "B", "C", "D", "E"]
        opt_vals = [str(opts.get(L) or "") for L in letters]
        ca = q.get("correct_answer") or {}
        rebuilt.append(
            {
                "question_r": q.get("question_number"),
                "question_hi": "",
                "options_hi": {},
                "solution_hi": "",
                "question_en": q.get("question") or "",
                "options_en": {
                    str(i + 1): opt_vals[i] for i in range(5) if opt_vals[i]
                },
                "solution_en": q.get("solution") or q.get("explanation") or "",
                "answer": ca.get("option") or ca.get("answer") or "",
                "difficulty_level": "medium",
            }
        )
    return questions_to_rows(rebuilt, set_name=set_name)


def save_docx(
    doc: dict[str, Any],
    path: Path,
    *,
    set_name: str | None = None,
    rows: list[dict[str, str]] | None = None,
) -> Path:
    """Write a full question-set Word file (Hindi + English + answer + solution)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    path.parent.mkdir(parents=True, exist_ok=True)
    meta = doc.get("document") or {}
    paper = set_name or meta.get("title") or meta.get("source_file") or "Paper Name"
    if rows is None:
        rows = _rows_from_doc(doc, set_name=str(paper))

    document = Document()
    title = document.add_heading("Solved Question Paper", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = document.add_paragraph()
    info.add_run("Set / Paper: ").bold = True
    info.add_run(str(paper))
    info = document.add_paragraph()
    info.add_run("Source: ").bold = True
    info.add_run(str(meta.get("source_file") or ""))
    info = document.add_paragraph()
    info.add_run("Total questions: ").bold = True
    info.add_run(str(len(rows)))
    info = document.add_paragraph()
    info.add_run("Generated: ").bold = True
    info.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    document.add_paragraph()

    for row in rows:
        qn = row.get("question_r") or ""
        document.add_heading(f"Question {qn}", level=1)

        q_hi = (row.get("question_hi") or "").strip()
        q_en = (row.get("question_en") or "").strip()

        if q_hi:
            p = document.add_paragraph()
            p.add_run("प्रश्न (Hindi):\n").bold = True
            p.add_run(q_hi)
        if q_en:
            p = document.add_paragraph()
            p.add_run("Question (English):\n").bold = True
            p.add_run(q_en)

        opts_hi = [
            row.get(f"option{i}_hi") or "" for i in range(1, 6)
        ]
        opts_en = [
            row.get(f"option{i}_en") or "" for i in range(1, 6)
        ]
        letters = ["A", "B", "C", "D", "E"]

        if any(opts_hi):
            document.add_paragraph().add_run("विकल्प (Hindi):").bold = True
            for i, text in enumerate(opts_hi):
                if text.strip():
                    document.add_paragraph(
                        f"{letters[i]}. {text}", style="List Bullet"
                    )

        if any(opts_en):
            document.add_paragraph().add_run("Options (English):").bold = True
            for i, text in enumerate(opts_en):
                if text.strip():
                    document.add_paragraph(
                        f"{letters[i]}. {text}", style="List Bullet"
                    )

        ans = (row.get("answer") or "").strip()
        p = document.add_paragraph()
        p.add_run("Correct Answer: ").bold = True
        p.add_run(ans)

        # Map letter answer to option text when possible
        if len(ans) == 1 and ans.upper() in letters:
            idx = letters.index(ans.upper())
            chosen = (opts_hi[idx] or opts_en[idx] or "").strip()
            if chosen:
                p = document.add_paragraph()
                p.add_run("Answer Text: ").bold = True
                p.add_run(chosen)

        sol_hi = _strip_html(row.get("solution_hi") or "")
        sol_en = _strip_html(row.get("solution_en") or "")
        if sol_hi:
            p = document.add_paragraph()
            p.add_run("हल (Hindi):\n").bold = True
            p.add_run(sol_hi)
        if sol_en:
            p = document.add_paragraph()
            p.add_run("Solution (English):\n").bold = True
            p.add_run(sol_en)

        diff = (row.get("difficulty_level") or "").strip()
        if diff:
            p = document.add_paragraph()
            p.add_run("Difficulty: ").bold = True
            p.add_run(diff)

        document.add_paragraph()

    try:
        document.save(path)
        return path
    except PermissionError:
        alt = path.with_name(
            f"{path.stem}_{datetime.now().strftime('%H%M%S')}{path.suffix}"
        )
        document.save(alt)
        print(f"WARNING: '{path.name}' locked. Saved as: {alt.name}", flush=True)
        return alt


def save_docx_from_csv(csv_path: Path, docx_path: Path, *, set_name: str | None = None) -> Path:
    csv_path = Path(csv_path)
    with csv_path.open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    # ensure header fields exist
    cleaned = []
    for row in rows:
        cleaned.append({h: row.get(h, "") for h in CSV_HEADERS})
    paper = set_name or (cleaned[0].get("set_name") if cleaned else "Paper Name")
    doc = {
        "document": {
            "title": paper,
            "source_file": csv_path.name,
            "source_type": "csv",
            "total_questions": len(cleaned),
        },
        "questions": [],
    }
    return save_docx(doc, Path(docx_path), set_name=paper, rows=cleaned)


def export_all(
    doc: dict[str, Any],
    out_dir: str | Path,
    *,
    set_name: str | None = None,
    questions_raw: list[dict[str, Any]] | None = None,
) -> dict[str, Path]:
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    if questions_raw is not None:
        doc = dict(doc)
        doc["questions_raw"] = questions_raw
    return {
        "json": save_json(doc, out / "solved_questions.json"),
        "docx": save_docx(doc, out / "solved_questions.docx", set_name=set_name),
        "report": save_report(doc, out / "processing_report.json"),
    }
